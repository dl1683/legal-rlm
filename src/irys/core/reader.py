"""Document reader for PDF, DOCX, and MHT files.

Extracts text with page/section preservation for citation tracking.
"""

from pathlib import Path
from dataclasses import dataclass
from typing import Optional
import re
import email
from email import policy
from html.parser import HTMLParser

import fitz  # PyMuPDF
from docx import Document


class HTMLTextExtractor(HTMLParser):
    """Extract text from HTML, stripping tags."""

    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.skip_tags = {'script', 'style', 'head', 'meta', 'link'}
        self.current_skip = 0

    def handle_starttag(self, tag, attrs):
        if tag.lower() in self.skip_tags:
            self.current_skip += 1
        # Add spacing for block elements
        if tag.lower() in {'p', 'div', 'br', 'li', 'tr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            self.text_parts.append('\n')

    def handle_endtag(self, tag):
        if tag.lower() in self.skip_tags:
            self.current_skip = max(0, self.current_skip - 1)
        if tag.lower() in {'p', 'div', 'li', 'tr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            self.text_parts.append('\n')

    def handle_data(self, data):
        if self.current_skip == 0:
            self.text_parts.append(data)

    def get_text(self) -> str:
        return ''.join(self.text_parts)


@dataclass
class PageContent:
    """Content from a single page."""
    page_num: int
    text: str


@dataclass
class DocumentContent:
    """Full document content with metadata."""
    path: str
    filename: str
    file_type: str
    page_count: int
    pages: list[PageContent]
    total_chars: int

    @property
    def full_text(self) -> str:
        """Get full document text with page markers."""
        parts = []
        for page in self.pages:
            parts.append(f"\n--- PAGE {page.page_num} ---\n")
            parts.append(page.text)
        return "".join(parts)

    def get_page_range(self, start: int, end: int) -> str:
        """Get text from a page range (1-indexed)."""
        parts = []
        for page in self.pages:
            if start <= page.page_num <= end:
                parts.append(f"\n--- PAGE {page.page_num} ---\n")
                parts.append(page.text)
        return "".join(parts)

    def get_excerpt(self, max_chars: int = 5000) -> str:
        """Get excerpt of document up to max_chars."""
        text = self.full_text
        if len(text) <= max_chars:
            return text
        return text[:max_chars] + f"\n\n[...truncated, {self.total_chars - max_chars} more chars...]"


class DocumentReader:
    """Read and extract text from PDF, DOCX, TXT, and MHT files.

    Note: Old .doc (binary) and .rtf formats are NOT supported.
    Convert to .docx or .pdf before processing.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".mht", ".mhtml"}

    def read(self, path: Path | str) -> DocumentContent:
        """Read a document and extract text."""
        path = Path(path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._read_pdf(path)
        elif suffix == ".docx":
            return self._read_docx(path)
        elif suffix == ".txt":
            return self._read_txt(path)
        elif suffix in {".mht", ".mhtml"}:
            return self._read_mht(path)
        elif suffix in {".doc", ".rtf"}:
            raise ValueError(
                f"Unsupported legacy format: {suffix}. "
                f"Please convert to .docx or .pdf first."
            )
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _read_pdf(self, path: Path) -> DocumentContent:
        """Extract text from PDF with page structure."""
        doc = fitz.open(path)
        pages = []
        total_chars = 0

        try:
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                text = self._clean_text(text)
                pages.append(PageContent(page_num=page_num, text=text))
                total_chars += len(text)
        finally:
            doc.close()

        return DocumentContent(
            path=str(path),
            filename=path.name,
            file_type="pdf",
            page_count=len(pages),
            pages=pages,
            total_chars=total_chars,
        )

    def _read_docx(self, path: Path) -> DocumentContent:
        """Extract text from DOCX."""
        doc = Document(path)

        # DOCX doesn't have real pages, treat as single page
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        text = self._clean_text(text)

        # Also extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            text += "\n\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n"

        pages = [PageContent(page_num=1, text=text)]

        return DocumentContent(
            path=str(path),
            filename=path.name,
            file_type="docx",
            page_count=1,
            pages=pages,
            total_chars=len(text),
        )

    def _read_txt(self, path: Path) -> DocumentContent:
        """Read plain text file."""
        text = path.read_text(encoding="utf-8", errors="replace")
        text = self._clean_text(text)

        pages = [PageContent(page_num=1, text=text)]

        return DocumentContent(
            path=str(path),
            filename=path.name,
            file_type="txt",
            page_count=1,
            pages=pages,
            total_chars=len(text),
        )

    def _read_mht(self, path: Path) -> DocumentContent:
        """Read MIME HTML (MHT/MHTML) file and extract text.

        MHT files are MIME-encoded web archives that contain HTML content.
        We extract the HTML and convert it to plain text.
        """
        # Read the raw content
        raw_content = path.read_bytes()

        # Try to parse as MIME message
        try:
            msg = email.message_from_bytes(raw_content, policy=policy.default)
            html_content = None

            # Walk through MIME parts to find HTML
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/html':
                        payload = part.get_payload(decode=True)
                        if payload:
                            # Try different encodings
                            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                                try:
                                    html_content = payload.decode(encoding)
                                    break
                                except UnicodeDecodeError:
                                    continue
                        break
            else:
                # Single-part message
                payload = msg.get_payload(decode=True)
                if payload:
                    for encoding in ['utf-8', 'latin-1', 'cp1252']:
                        try:
                            html_content = payload.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue

            # Fallback: try reading as raw text with HTML
            if not html_content:
                raw_text = raw_content.decode('utf-8', errors='replace')
                # Look for HTML content between markers
                if '<html' in raw_text.lower():
                    html_content = raw_text
        except Exception:
            # Final fallback: read as text and extract what we can
            html_content = path.read_text(encoding='utf-8', errors='replace')

        # Extract text from HTML
        if html_content:
            # Handle quoted-printable encoding artifacts
            html_content = html_content.replace('=\n', '')  # Line continuations
            html_content = re.sub(r'=([0-9A-Fa-f]{2})',
                                  lambda m: chr(int(m.group(1), 16)),
                                  html_content)

            # Use regex-based extraction (more robust than HTMLParser for MHT files)
            # Remove script and style content
            text = re.sub(r'<script[^>]*>.*?</script>', '', html_content,
                          flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text,
                          flags=re.DOTALL | re.IGNORECASE)
            # Strip remaining tags
            text = re.sub(r'<[^>]+>', ' ', text)
        else:
            text = ""

        text = self._clean_text(text)
        pages = [PageContent(page_num=1, text=text)]

        return DocumentContent(
            path=str(path),
            filename=path.name,
            file_type="mht",
            page_count=1,
            pages=pages,
            total_chars=len(text),
        )

    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Normalize horizontal whitespace (preserve newlines)
        text = re.sub(r'[^\S\n]+', ' ', text)
        # Clean up spaces around newlines
        text = re.sub(r' ?\n ?', '\n', text)
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def get_page_count(self, path: Path | str) -> int:
        """Get page count without reading full content."""
        path = Path(path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            doc = fitz.open(path)
            count = len(doc)
            doc.close()
            return count
        else:
            return 1  # Non-paginated formats

    def can_read(self, path: Path | str) -> bool:
        """Check if file type is supported."""
        path = Path(path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS
